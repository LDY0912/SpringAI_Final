from pathlib import Path
from urllib.parse import unquote, urlparse
import unicodedata

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import html


ROOT = Path(__file__).resolve().parents[1]
HTML_PATH = ROOT / "SKALA_HelpDesk_AI_종합실습_보고서.html"
DOCX_PATH = ROOT / "SKALA_HelpDesk_AI_종합실습_보고서.docx"

NAVY = "10203C"
BLUE = "173E79"
PINK = "EF2455"
MUTED = "667792"
LIGHT_BLUE = "EAF0F8"
LIGHT_GRAY = "F5F7FA"
CODE_BG = "111B2E"
CODE_FG = "EDF4FF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_false(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("- ")
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    paragraph.add_run(" -")


def configure_document(document):
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Apple SD Gothic Neo"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = RGBColor.from_string("17233A")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color in (
        ("Title", 28, NAVY),
        ("Heading 1", 18, NAVY),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, BLUE),
    ):
        style = document.styles[style_name]
        style.font.name = "Apple SD Gothic Neo"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    document.styles["Heading 1"].paragraph_format.space_before = Pt(16)
    document.styles["Heading 1"].paragraph_format.space_after = Pt(8)
    document.styles["Heading 2"].paragraph_format.space_before = Pt(11)
    document.styles["Heading 2"].paragraph_format.space_after = Pt(5)
    document.styles["Heading 3"].paragraph_format.space_before = Pt(9)
    document.styles["Heading 3"].paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "SKALA · Spring AI Final Lab"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_heading(document, text, level, page_break=False):
    paragraph = document.add_paragraph(style="Title" if level == 0 else f"Heading {level}")
    if page_break:
        paragraph.paragraph_format.page_break_before = True
    paragraph.add_run(text.strip())
    if level == 1:
        p_pr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "5")
        bottom.set(qn("w:color"), PINK)
        borders.append(bottom)
        p_pr.append(borders)
    return paragraph


def add_code(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.25)
    paragraph.paragraph_format.right_indent = Cm(0.25)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.0
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CODE_BG)
    p_pr.append(shd)
    run = paragraph.add_run(text.strip())
    run.font.name = "Menlo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    run.font.size = Pt(7.4)
    run.font.color.rgb = RGBColor.from_string(CODE_FG)


def add_result(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.25)
    paragraph.paragraph_format.right_indent = Cm(0.25)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.0
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    run = paragraph.add_run(text.strip())
    run.font.name = "Menlo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    run.font.size = Pt(7.4)
    run.font.color.rgb = RGBColor.from_string("23334E")


def add_box(document, text, kind):
    colors = {
        "info": ("EEF5FF", "2457A6"),
        "note": ("FFF8E6", "E2A316"),
        "danger": ("FFF0F3", PINK),
        "summary": ("F2F6FB", PINK),
        "meta": ("F2F6FB", BLUE),
    }
    fill, accent = colors.get(kind, ("F5F7FA", BLUE))
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 110, 150, 110, 150)
    paragraph = cell.paragraphs[0]
    paragraph.add_run(text.strip())
    left_border = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "22")
    left.set(qn("w:color"), accent)
    left_border.append(left)
    cell._tc.get_or_add_tcPr().append(left_border)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_html_table(document, node):
    rows = node.xpath("./tr|./thead/tr|./tbody/tr")
    if not rows:
        return
    width = max(len(row.xpath("./th|./td")) for row in rows)
    table = document.add_table(rows=0, cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, source_row in enumerate(rows):
        cells = source_row.xpath("./th|./td")
        row = table.add_row()
        set_repeat_false(row)
        if row_index == 0:
            set_repeat_table_header(row)
        for column_index, source_cell in enumerate(cells):
            target = row.cells[column_index]
            target.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(target)
            target.text = " ".join(source_cell.text_content().split())
            for paragraph in target.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.2)
                    if source_cell.tag == "th" or row_index == 0:
                        run.font.bold = True
            if source_cell.tag == "th" or row_index == 0:
                set_cell_shading(target, LIGHT_BLUE)
            if "pass" in (source_cell.get("class") or ""):
                for run in target.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(22, 120, 71)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(document, node):
    source = node.get("src", "")
    if source.startswith("file://"):
        path = Path(unquote(urlparse(source).path))
    else:
        path = (HTML_PATH.parent / source).resolve()
    if not path.exists():
        document.add_paragraph(f"[이미지 누락: {path}]")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.5))


def has_block_children(node):
    return any(child.tag in {"p", "h1", "h2", "h3", "pre", "table", "ul", "img", "section", "div"} for child in node)


def render_node(document, node):
    tag = node.tag.lower() if isinstance(node.tag, str) else ""
    classes = set((node.get("class") or "").split())
    text = node.text_content().strip()

    if tag in {"style", "script"}:
        return
    if tag == "h1":
        add_heading(document, text, 0)
    elif tag == "h2":
        add_heading(document, text, 1, "page-break" in classes)
    elif tag == "h3":
        add_heading(document, text, 2)
    elif tag == "p":
        paragraph = document.add_paragraph(text)
        if "subtitle" in classes:
            for run in paragraph.runs:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor.from_string(MUTED)
        if "caption" in classes:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.italic = True
                run.font.color.rgb = RGBColor.from_string(MUTED)
        if "small" in classes:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor.from_string(MUTED)
        if "pass" in classes:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(22, 120, 71)
    elif tag == "pre":
        if "result" in classes:
            add_result(document, text)
        else:
            add_code(document, text)
    elif tag == "table":
        add_html_table(document, node)
    elif tag == "ul":
        for item in node.xpath("./li"):
            document.add_paragraph(item.text_content().strip(), style="List Bullet")
    elif tag == "img":
        add_image(document, node)
    elif tag in {"div", "section", "body"}:
        if tag == "div" and classes.intersection({"info", "note", "danger"}) and not has_block_children(node):
            add_box(document, text, next(iter(classes.intersection({"info", "note", "danger"}))))
        elif tag == "div" and classes.intersection({"summary", "meta"}):
            for child in node:
                render_node(document, child)
        elif has_block_children(node):
            for child in node:
                render_node(document, child)
        elif text:
            document.add_paragraph(text)


def main():
    source = unicodedata.normalize("NFC", HTML_PATH.read_text(encoding="utf-8"))
    corrections = {
        "건증 · 실시간": "검증 · 실시간",
        "검젱일": "검증일",
        "\u1111로젝트": "프로젝트",
        "스\u1110림밍·옵점대시보드가 화긴되었다": "스트리밍·운영 대시보드가 확인되었다",
        "Golden Set 품질 평가눈 실습 범위에서 의도적우로 제외했다": "Golden Set 품질 평가는 실습 범위에서 의도적으로 제외했다",
        "1. 실습 목표\u1119 구성": "1. 실습 목표와 구성",
        "사내 구정 문서의 근거로 답하고, 주문·티켓에 실시간 데이터룰 연동하며, 앞 대화의 맥락을 유지하고, 교환·환불은 사람의 숭인 절차를 거치는 것이다": "사내 규정 문서의 근거로 답하고, 주문·티켓에 실시간 데이터를 연동하며, 앞 대화의 맥락을 유지하고, 교환·환불은 사람의 승인 절차를 거치는 것이다",
        "에플리케이션": "애플리케이션",
        "이름 · 유해 · 인제스트 · 임베딩 · 저장": "웹 UI · ChatClient · Advisor · Repository",
        "프롬\u1111트": "프롬프트",
        "옵고개 상담 UI": "웹고객 상담 UI",
        "웹고객 상담 UI": "웹 고객 상담 UI",
        "별도 인\u1111라 없이 실행 가능한 SimpleVectorStore랄 사용했다. 재시작 시 재인제스트가 피요하다": "별도 인프라 없이 실행 가능한 SimpleVectorStore를 사용했다. 재시작 시 재인제스트가 필요하다",
        "2.3 공통 호출이 연결된 겆역 주이아항": "2.3 공통 호출 및 보안 주의사항",
        "아래 명령이나 비밀번호느 출력에 \u1111포시하지 않았다": "API 키는 출력에 표시하지 않았으며 실습용 기본 비밀번호만 명령에 사용했다",
        "사내 구정 문서 재새인": "사내 규정 문서 재색인",
        "기존 청크랄": "기존 청크를",
        "메타데이터랄": "메타데이터를",
        "인제스트 품질 화긴": "인제스트 품질 확인",
        "결과물을 화긴했다": "결과물을 확인했다",
        "출처 모두 반환다": "출처를 모두 반환한다",
        "4.2 가튼 세션": "4.2 같은 세션",
        "현재 배송 중이며로, 반품은 상품을 받은 날부터 7일 이내에 신청해야 가능합니다": "현재 배송 중이므로, 반품은 상품을 받은 날부터 7일 이내에 신청할 수 있습니다",
        "4.3 티켓 생성와 사람 승인": "4.3 티켓 생성과 사람 승인",
        "5.1 프롬프트 인젝션 차단이 메모리 아\u1111 동작": "5.1 프롬프트 인젝션 차단이 메모리 앞에서 동작",
        "중간 토큰 \u1109트림": "중간 토큰 스트림",
        "별도 이벤트로 반환다": "별도 이벤트로 반환한다",
        "운영자가 질문할 수 있다": "운영자가 확인할 수 있다",
        "6. 옵 기반 결과 화면": "6. 웹 기반 결과 화면",
        "고개 상답 UI": "고객 상담 UI",
        "옵 기능은 가산점 항목이다. 초틈간 기능 요구사항은 REST API·Spring AI 응답이며, 옵은 같은 기능을 브라우저 위한 추가 구현한 가산점 구현이다": "웹 기능은 가산점 항목이다. 핵심 기능 요구사항은 REST API·Spring AI 응답이며, 웹은 같은 기능을 브라우저에서도 확인할 수 있도록 추가한 구현이다",
        "7. 실습 자료 요구사항 대응\u1111\u11191": "7. 실습 자료 요구사항 대응표",
        "문서 재새인·메타데이터·품질 화긴": "문서 재색인·메타데이터·품질 확인",
        "SafetyAdvisor를 Memory 아\u1111에 배치": "SafetyAdvisor를 Memory 앞에 배치",
        "오로 시·타임아웃 시·60초 폴백 응답": "오류·타임아웃 시 60초 폴백 응답",
        "가산점 옵 UI": "가산점 웹 UI",
        "본 \u1111로젝트는 사내 구정 지식": "본 프로젝트는 사내 규정 지식",
        "실습 이행 여부를 낮추었다": "실습 이행 여부를 쉽게 확인할 수 있게 했다",
        "이번 제출하여 본 실습에서는 제외했다. 단위, 자동화 테스트와 실제 호출 결과로 현재 점검을 묭\u1111나다": "이번 제출에서는 제외했다. 단위·자동화 테스트와 실제 호출 결과로 현재 상태를 점검했다",
        "실제로 화긴한 결과": "실제로 확인한 결과",
        "옵": "웹",
        "고개": "고객",
        "상답": "상담",
        "구정": "규정",
        "인\u1111라": "인프라",
        "SimpleVectorStore랄": "SimpleVectorStore를",
        "피요": "필요",
        "초틈간": "핵심",
        "브라우저 위한": "브라우저에서도 확인할 수 있도록",
        "추가 구현한 가산점 구현": "추가한 가산점 구현",
        "진행한 되": "진행한 뒤",
        "그거 어떻게 돼어요?": "그거 어떻게 됐어요?",
        "화긴": "확인",
        "가튼": "같은",
        "䍑\u1173\u11ba인다": "뜻이다",
        "SimpleVectorStore</code>랄": "SimpleVectorStore</code>를",
        "동기 JSON/SSE 선택": "동기 JSON 고정",
        "시나리오와 JSON/SSE 선택": "시나리오와 동기 JSON 응답",
    }
    for wrong, right in corrections.items():
        source = source.replace(wrong, right)
    tree = html.fromstring(source)
    body = tree.find("body")
    document = Document()
    configure_document(document)
    render_node(document, body)

    core = document.core_properties
    core.title = "SKALA HelpDesk AI 종합 실습 보고서"
    core.subject = "Spring AI RAG, Tool, Memory, Safety, SSE 종합 실습 검증"
    core.author = "SKALA 종합 실습"
    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
